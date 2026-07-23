"""Build the Gemini handoff Word document from its Markdown source."""

from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


def set_run_font(run, name: str, size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def add_inline_markdown(paragraph, text: str) -> None:
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, "Consolas", 9.5)
            run.font.color.rgb = RGBColor(44, 62, 80)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, "Microsoft YaHei", 10.5)
            run.bold = True
        else:
            run = paragraph.add_run(part)
            set_run_font(run, "Microsoft YaHei", 10.5)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.35

    for level, size, color in [
        (1, 18, "16324F"),
        (2, 14, "1F4E78"),
        (3, 12, "2F5597"),
    ]:
        style = document.styles[f"Heading {level}"]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.add_run("mutiAI Gemini 前端交接说明  |  ")
    add_page_number(footer)


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(120)
    title = paragraph.add_run("mutiAI")
    set_run_font(title, "Microsoft YaHei", 28)
    title.bold = True
    title.font.color.rgb = RGBColor(22, 50, 79)

    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = subtitle_p.add_run("Gemini 前端项目交接说明")
    set_run_font(subtitle, "Microsoft YaHei", 20)
    subtitle.font.color.rgb = RGBColor(31, 78, 120)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(30)
    meta_run = meta.add_run(
        f"文档版本：V1 bootstrap\n生成日期：{date.today().isoformat()}\n"
        "核心仓库：https://github.com/Purewo/mutiAI\n"
        "前端仓库：https://github.com/Purewo/mutiAI-Gemini"
    )
    set_run_font(meta_run, "Microsoft YaHei", 10.5)
    meta_run.font.color.rgb = RGBColor(89, 89, 89)

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(40)
    note_run = note.add_run("用于 Gemini 在 Google AI Studio 中快速接管前端工作的战前资料")
    set_run_font(note_run, "Microsoft YaHei", 10.5)
    note_run.italic = True

    document.add_page_break()


def add_code_block(document: Document, lines: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(0.5)
    paragraph.paragraph_format.right_indent = Cm(0.5)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F6F8")
    paragraph._p.get_or_add_pPr().append(shading)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, "Consolas", 9)
    run.font.color.rgb = RGBColor(44, 62, 80)


def render_markdown(document: Document, source: str) -> None:
    in_code = False
    code_lines: list[str] = []

    for raw_line in source.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            if level == 1:
                continue
            paragraph = document.add_paragraph(style=f"Heading {level - 1}")
            add_inline_markdown(paragraph, heading.group(2))
            continue

        bullet = re.match(r"^-\s+(.+)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_markdown(paragraph, bullet.group(1))
            continue

        numbered = re.match(r"^\d+\.\s+(.+)$", line)
        if numbered:
            paragraph = document.add_paragraph(style="List Number")
            add_inline_markdown(paragraph, numbered.group(1))
            continue

        quote = re.match(r"^>\s*(.+)$", line)
        if quote:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.8)
            run = paragraph.add_run(quote.group(1))
            set_run_font(run, "Microsoft YaHei", 10.5)
            run.italic = True
            run.font.color.rgb = RGBColor(89, 89, 89)
            continue

        paragraph = document.add_paragraph()
        add_inline_markdown(paragraph, line)


def main() -> int:
    if len(sys.argv) < 3:
        raise SystemExit(
            "usage: build_handoff_docx.py SOURCE.md OUTPUT.docx [OUTPUT.docx ...]"
        )

    source_path = Path(sys.argv[1]).resolve()
    outputs = [Path(item).resolve() for item in sys.argv[2:]]
    source = source_path.read_text(encoding="utf-8")

    document = Document()
    configure_document(document)
    add_title_page(document)
    render_markdown(document, source)

    for index, output in enumerate(outputs):
        output.parent.mkdir(parents=True, exist_ok=True)
        document.save(output)
        print(f"created: {output}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
